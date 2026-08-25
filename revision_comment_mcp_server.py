"""
revision_comment_mcp_server.py
----------------------------------
도면 리비전(개정판) 변경사항 비교 + 발주처 Comment Sheet 답변 초안 자동화를 위한
MCP 서버입니다. 아래 두 업무를 하나의 흐름으로 연결합니다.

  A) 도면이 개정될 때, 이전 리비전 대비 어디가 바뀌었는지 자동으로 찾아준다.
  B) 발주처가 과거에 준 코멘트와 비슷한 코멘트가 다시 왔을 때, 과거 회신 이력에서
     유사 사례를 찾아주고, 그 위에 최신 개정 내용을 반영한 답변 초안(Word)을 만든다.

이 서버(파이썬 코드)가 하는 일은 "정확한 계산·검색·파일 처리"까지입니다.
- 어디가 실질적으로 바뀐 것인지(단순 노이즈 vs 의미있는 변경) 최종 판단
- 어떤 과거 사례가 지금 코멘트와 진짜 비슷한지 최종 판단
- 사내 어투에 맞춰 완결된 답변 문장으로 다듬는 것
은 이 tool을 호출하는 AI(Claude 등)가 담당하고, 최종 방향 지시와 검수는 사람(담당자)이 합니다.
서버는 후보를 찾아 넘겨주고, AI/사람이 정한 결과를 정해진 문서 양식대로 기록하는 역할만 합니다.

제공 기능:
  1) 리비전 도면 비교 (diff_drawing_revisions) — 페이지를 이미지로 렌더링해 격자 단위로
     변경 가능성이 높은 영역을 표시. 실제 "무엇이 바뀌었는지"의 의미 해석은 AI가 이미지를 보고 판단.
  2) 과거 Comment Sheet 파싱 (parse_comment_sheet)
  3) 유사 코멘트/과거 답변 검색 (search_similar_comments)
  4) Comment Sheet 답변 초안 Word 파일 생성 (create_comment_sheet_draft)

실행 방법:
  1) pip install mcp pymupdf pillow numpy python-docx
  2) python revision_comment_mcp_server.py   (단독 실행 시 stdio로 대기)
  3) Claude Desktop 설정(claude_desktop_config.json)에 등록:
     {
       "mcpServers": {
         "revision-comment-tools": {
           "command": "python",
           "args": ["/절대/경로/revision_comment_mcp_server.py"]
         }
       }
     }

주의:
  - COMMENT_SHEET_COLUMN_KEYWORDS, 표 구조 가정은 예시입니다.
    실제 발주처 Comment Sheet 양식(컬럼 구성)에 맞춰 반드시 확인/수정하세요.
  - 유사도 검색은 difflib 기반 문자열 유사도(개발 단계 MVP)입니다.
    실제 사용 시 오탐/누락이 있을 수 있어, 검색 결과는 AI가 다시 읽고 재판단하는 걸 전제로 합니다.
"""

import io
from datetime import date
from difflib import SequenceMatcher
from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
from PIL import Image as PILImage, ImageDraw
from docx import Document
from mcp.server.fastmcp import FastMCP, Image

mcp = FastMCP("revision-comment-tools")

BASE_DIR = Path(__file__).parent
HISTORY_DIR = BASE_DIR / "history"
OUTPUT_DIR = BASE_DIR / "output"
HISTORY_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Comment Sheet 표에서 컬럼 역할을 찾을 때 쓰는 헤더 키워드 (예시 - 실제 양식에 맞게 수정)
# "no"는 "번호"만 단독으로 쓰인 경우만 항목번호로 보고, "도면번호"처럼 복합어에 포함된
# 경우는 drawing_no로 잡히도록 exact 매칭을 쓴다 (아래 _match_column 참고).
COMMENT_SHEET_COLUMN_EXACT = {"no": "no", "no.": "no", "번호": "no"}
COMMENT_SHEET_COLUMN_KEYWORDS = {
    "drawing_no": ["도면", "drawing", "도서", "항목"],
    "comment": ["코멘트", "comment", "검토의견", "의견"],
    "response": ["회신", "답변", "response", "reply", "반영"],
    "status": ["상태", "status", "구분"],
}


def _match_column(header_text: str) -> str | None:
    text = header_text.strip().lower()
    if text in COMMENT_SHEET_COLUMN_EXACT:
        return COMMENT_SHEET_COLUMN_EXACT[text]
    for role, keywords in COMMENT_SHEET_COLUMN_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            return role
    return None


def _parse_comment_table(table) -> list[dict]:
    """docx 표 하나를 파싱해 role별 컬럼 값을 담은 행 리스트로 변환."""
    if not table.rows:
        return []

    header_cells = [c.text for c in table.rows[0].cells]
    col_roles = [_match_column(h) for h in header_cells]

    # 헤더에서 역할을 하나도 못 찾으면, 자리(위치) 기반으로 추정 (No/구분/코멘트/회신 순 가정)
    if not any(col_roles):
        fallback_roles = ["no", "drawing_no", "comment", "response", "status"]
        col_roles = [fallback_roles[i] if i < len(fallback_roles) else None for i in range(len(header_cells))]

    rows = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        record = {"raw": cells}
        for role, value in zip(col_roles, cells):
            if role:
                record[role] = value
        if any(v for k, v in record.items() if k != "raw"):
            rows.append(record)
    return rows


@mcp.tool()
def parse_comment_sheet(docx_path: str) -> dict:
    """
    발주처 Comment Sheet(.docx) 파일을 열어 표를 파싱하고,
    No/도면번호/코멘트/회신/상태 컬럼으로 정리된 행 리스트를 돌려준다.
    표 헤더에서 역할을 못 찾으면 위치 기반(No, 구분, 코멘트, 회신, 상태 순)으로 추정한다.
    """
    path = Path(docx_path)
    if not path.exists():
        return {"ok": False, "error": f"파일을 찾을 수 없습니다: {docx_path}"}

    try:
        doc = Document(path)
    except Exception as e:
        return {"ok": False, "error": f"docx를 열 수 없습니다: {e}"}

    all_rows = []
    for table in doc.tables:
        all_rows.extend(_parse_comment_table(table))

    return {"ok": True, "source": str(path), "row_count": len(all_rows), "rows": all_rows}


@mcp.tool()
def search_similar_comments(comment_text: str, history_dir: str = "", top_k: int = 5) -> dict:
    """
    history_dir(기본값: 이 서버의 history/ 폴더) 안의 모든 Comment Sheet(.docx)를 훑어,
    입력한 comment_text와 문자열 유사도가 높은 과거 코멘트+회신 사례를 top_k개 찾아준다.

    주의: 이 유사도는 difflib 기반 표면적 문자열 비교라서, 진짜 의미상 유사한지는
    반환된 후보를 직접 읽고 AI가 다시 판단해야 한다. 결과가 비어있으면 유사 사례가
    없다는 뜻이므로, 코멘트의 핵심 요지를 새로 파악해 답변을 작성해야 한다.
    """
    directory = Path(history_dir) if history_dir else HISTORY_DIR
    if not directory.exists():
        return {"ok": False, "error": f"디렉토리를 찾을 수 없습니다: {directory}"}

    docx_files = sorted(directory.rglob("*.docx"))
    if not docx_files:
        return {"ok": True, "candidates": [], "note": f"{directory} 에 .docx 이력 파일이 없습니다."}

    query = comment_text.strip().lower()
    candidates = []
    for f in docx_files:
        parsed = parse_comment_sheet(str(f))
        if not parsed.get("ok"):
            continue
        for row in parsed["rows"]:
            past_comment = row.get("comment", "")
            if not past_comment:
                continue
            score = SequenceMatcher(None, query, past_comment.strip().lower()).ratio()
            candidates.append({
                "score": round(score, 3),
                "source_file": f.name,
                "past_comment": past_comment,
                "past_response": row.get("response", ""),
                "drawing_no": row.get("drawing_no", ""),
                "status": row.get("status", ""),
            })

    candidates.sort(key=lambda c: c["score"], reverse=True)
    return {"ok": True, "candidates": candidates[:top_k], "searched_files": len(docx_files)}


@mcp.tool()
def create_comment_sheet_draft(drawing_number: str, revision: str, rows: list[dict], output_name: str = "") -> dict:
    """
    Comment Sheet 답변 초안 Word 파일을 생성한다.
    rows: [{"no": 1, "drawing_no": "...", "comment": "...", "response": "...", "status": "초안 - 검수 필요"}, ...]
    response는 AI가 search_similar_comments 결과와 diff_drawing_revisions로 파악한 개정 내용을
    바탕으로 작성한 답변 초안이어야 하며, 최종 제출 전 담당자 검수가 필요하다는 문구를 status에 남긴다.
    """
    if not rows:
        return {"ok": False, "error": "rows가 비어 있습니다."}

    filename = output_name or f"{drawing_number}_Rev{revision}_CommentSheet_Draft.docx"
    path = OUTPUT_DIR / filename

    doc = Document()
    doc.add_heading(f"Comment Sheet 답변 초안 - {drawing_number} (Rev.{revision})", level=1)
    doc.add_paragraph(f"작성일: {date.today().isoformat()}  |  ※ AI 초안 - 제출 전 담당자 최종 검수 필요")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ["No", "도면번호", "발주처 코멘트", "회신(초안)", "상태"]):
        cell.text = text

    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(row.get("no", i))
        cells[1].text = str(row.get("drawing_no", drawing_number))
        cells[2].text = str(row.get("comment", ""))
        cells[3].text = str(row.get("response", ""))
        cells[4].text = str(row.get("status", "초안 - 검수 필요"))

    doc.save(path)
    return {"ok": True, "path": str(path.resolve()), "filename": filename, "row_count": len(rows)}


def _render_page_gray(doc: fitz.Document, page_no_0based: int, dpi: int) -> tuple[PILImage.Image, bytes]:
    page = doc.load_page(page_no_0based)
    zoom = dpi / 72
    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    png_bytes = pixmap.tobytes("png")
    img = PILImage.open(io.BytesIO(png_bytes)).convert("RGB")
    return img, png_bytes


@mcp.tool()
def diff_drawing_revisions(
    old_pdf_path: str,
    new_pdf_path: str,
    page_pairs: list[list[int]] | None = None,
    dpi: int = 150,
    grid_rows: int = 20,
    grid_cols: int = 28,
    threshold: float = 0.012,
):
    """
    같은 도면의 이전 리비전(old_pdf_path)과 새 리비전(new_pdf_path) PDF를 페이지 단위로
    비교해, 격자(grid_rows x grid_cols)로 나눈 셀 중 변화가 감지된 영역을 새 리비전 이미지
    위에 빨간 박스로 표시해서 반환한다.

    page_pairs를 생략하면 같은 순서의 페이지끼리(1-1, 2-2, ...) 비교한다.
    예: [[1, 1], [3, 2]] 는 old의 1페이지-new의 1페이지, old의 3페이지-new의 2페이지를 비교.

    이 tool은 "픽셀 단위로 달라진 영역이 어디인지"만 찾아줄 뿐, 그게 치수 변경인지 심볼
    추가인지 등 의미 해석은 반환된 이미지를 직접 보고 판단해야 한다.

    threshold/grid 기본값은 "두 PDF가 같은 프로그램에서 벡터로 뽑은, 래스터화 잡음이
    거의 없는 파일"이라는 가정하에 얇은 선(도면 특유의 가는 선/문자) 변화까지 잡히도록
    낮게 잡은 값이다. 스캔본처럼 페이지마다 미세하게 흔들리는 PDF를 비교하면 배경 노이즈
    때문에 오탐(실제로는 안 바뀐 영역이 감지됨)이 늘어날 수 있으니, 그 경우 threshold를
    0.03~0.05 수준으로 올려서 재시도해야 한다.
    """
    old_path, new_path = Path(old_pdf_path), Path(new_pdf_path)
    if not old_path.exists():
        return [f"이전 리비전 파일을 찾을 수 없습니다: {old_pdf_path}"]
    if not new_path.exists():
        return [f"새 리비전 파일을 찾을 수 없습니다: {new_pdf_path}"]

    try:
        old_doc = fitz.open(old_path)
        new_doc = fitz.open(new_path)
    except Exception as e:
        return [f"PDF를 열 수 없습니다: {e}"]

    if page_pairs is None:
        n = min(old_doc.page_count, new_doc.page_count)
        page_pairs = [[i, i] for i in range(1, n + 1)]

    result: list = [
        f"'{old_path.name}'(전체 {old_doc.page_count}p) vs '{new_path.name}'(전체 {new_doc.page_count}p) "
        f"— {len(page_pairs)}개 페이지쌍 비교 (격자 {grid_rows}x{grid_cols}, threshold={threshold})"
    ]

    for old_no, new_no in page_pairs:
        if not (1 <= old_no <= old_doc.page_count) or not (1 <= new_no <= new_doc.page_count):
            result.append(f"--- 잘못된 페이지 번호: old={old_no}, new={new_no} (건너뜀) ---")
            continue

        old_img, _ = _render_page_gray(old_doc, old_no - 1, dpi)
        new_img, _ = _render_page_gray(new_doc, new_no - 1, dpi)

        # 두 이미지 크기가 다르면 old를 new 크기에 맞춰 리사이즈 (용지/스케일 차이 보정)
        if old_img.size != new_img.size:
            old_img = old_img.resize(new_img.size, PILImage.BILINEAR)

        old_arr = np.asarray(old_img.convert("L"), dtype=np.float32) / 255.0
        new_arr = np.asarray(new_img.convert("L"), dtype=np.float32) / 255.0

        h, w = new_arr.shape
        cell_h, cell_w = h / grid_rows, w / grid_cols

        changed_cells = []
        for r in range(grid_rows):
            for c in range(grid_cols):
                y0, y1 = int(r * cell_h), int((r + 1) * cell_h)
                x0, x1 = int(c * cell_w), int((c + 1) * cell_w)
                cell_diff = np.abs(old_arr[y0:y1, x0:x1] - new_arr[y0:y1, x0:x1]).mean()
                if cell_diff > threshold:
                    changed_cells.append((r, c, round(float(cell_diff), 3)))

        annotated = new_img.copy()
        draw = ImageDraw.Draw(annotated)
        for r, c, _score in changed_cells:
            y0, y1 = int(r * cell_h), int((r + 1) * cell_h)
            x0, x1 = int(c * cell_w), int((c + 1) * cell_w)
            draw.rectangle([x0, y0, x1, y1], outline=(255, 0, 0), width=3)

        buf = io.BytesIO()
        annotated.save(buf, format="PNG")

        result.append(f"--- old p.{old_no} vs new p.{new_no}: 변경 의심 셀 {len(changed_cells)}/{grid_rows * grid_cols} ---")
        result.append(Image(data=buf.getvalue(), format="png"))

    old_doc.close()
    new_doc.close()
    return result


if __name__ == "__main__":
    mcp.run()
